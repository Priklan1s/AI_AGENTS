import streamlit as st
from graph.workflow import graph
from parsing.pdf_parser import load_pdf_text
from docx import Document
from knowledge_base.loader import load_knowledge
from knowledge_base.vector_store import create_vector_store
from langchain_ollama import OllamaLLM
import pandas as pd
from io import StringIO

st.title("AI QA Assistant")

# -----------------------------
# LLM для QA чата
# -----------------------------

chat_llm = OllamaLLM(model="qwen2.5:14b")

# -----------------------------
# PROMPT для генерации CSV тест-кейсов
# -----------------------------

ZEPHYR_CSV_PROMPT = """
Ты Senior QA Engineer.

Сгенерируй тест-кейсы на основе требований.

ВАЖНО:
Верни результат СТРОГО в формате CSV.

ПРАВИЛА:

1. НЕ используй markdown
2. НЕ добавляй пояснений
3. НЕ добавляй текст до или после CSV
4. Первая строка должна быть заголовком
5. Разделитель — запятая
6. Steps обязательно в кавычках
7. Каждая строка = один тест-кейс
8. Минимум 10 тест-кейсов
9. Используй английские заголовки колонок

ФОРМАТ CSV:

ID,Name,Objective,Preconditions,Steps,Expected Result,Traceability

ПРИМЕР:

TC01,Edit BOM item,Verify editing BOM item,BOM opened,"1. Select item; 2. Change attribute; 3. Save",Item updated successfully,REQ-01
TC02,Cancel editing,Verify cancel editing,BOM opened,"1. Change attribute; 2. Press cancel",Changes not saved,REQ-02

Сгенерируй тест-кейсы для требований ниже.
"""

# -----------------------------
# состояние чата
# -----------------------------

if "chat_history" not in st.session_state:
    st.session_state.chat_history = []

if "last_result" not in st.session_state:
    st.session_state.last_result = None


# -----------------------------
# загрузка требований
# -----------------------------

uploaded_file = st.file_uploader(
    "Загрузить требования",
    type=["pdf", "docx", "txt"]
)

requirements_text = ""

if uploaded_file:

    if uploaded_file.name.endswith(".pdf"):
        requirements_text = load_pdf_text(uploaded_file)

    else:
        requirements_text = uploaded_file.read().decode("utf-8")


# -----------------------------
# загрузка knowledge base
# -----------------------------

@st.cache_resource
def load_vector_store():

    knowledge_texts = load_knowledge("knowledge")

    return create_vector_store(
        knowledge_texts
    )


vector_store = load_vector_store()


# -----------------------------
# запуск AI pipeline
# -----------------------------

if st.button("Анализировать"):

    if not requirements_text:
        st.warning("Добавьте требования")
        st.stop()

    with st.spinner("AI анализирует требования..."):

        result = graph.invoke({
            "requirements": ZEPHYR_CSV_PROMPT + "\n\n" + requirements_text,
            "vector_store": vector_store
        })

    if result is None:
        st.error("AI pipeline вернул пустой результат")
        st.stop()

    st.session_state.last_result = result

    # -----------------------------
    # DEBUG
    # -----------------------------

    st.subheader("DEBUG RESULT")
    st.json(result)

    # -----------------------------
    # Анализ требований
    # -----------------------------

    st.subheader("Анализ требований")

    analysis = result.get(
        "analysis",
        "Анализ требований не был сгенерирован"
    )

    st.write(analysis)

    # -----------------------------
    # Риски
    # -----------------------------

    st.subheader("Риски")

    risks = result.get(
        "risks",
        "Риски не были сгенерированы"
    )

    st.write(risks)

    # -----------------------------
    # Тест-дизайн
    # -----------------------------

    st.subheader("Тест-дизайн")

    test_design = result.get(
        "test_design",
        "Тест-дизайн отсутствует"
    )

    st.write(test_design)

    # -----------------------------
    # Тест-кейсы
    # -----------------------------

    st.subheader("Тест-кейсы")

    testcases = result.get(
        "testcases",
        "Тест-кейсы не были созданы"
    )

    # очистка markdown
    testcases = testcases.replace("```csv", "")
    testcases = testcases.replace("```", "")
    testcases = testcases.strip()

    st.code(testcases)

    # -----------------------------
    # CSV экспорт для Zephyr
    # -----------------------------

    try:

        csv_buffer = StringIO(testcases)

        df = pd.read_csv(
            csv_buffer,
            quotechar='"',
            skipinitialspace=True
        )

        st.subheader("Таблица тест-кейсов")

        st.dataframe(df)

        csv_data = df.to_csv(index=False)

        st.download_button(
            label="Скачать CSV для Zephyr",
            data=csv_data,
            file_name="zephyr_testcases.csv",
            mime="text/csv"
        )

    except:

        st.warning("AI не вернул корректный CSV формат тест-кейсов")

    # -----------------------------
    # Генерация DOCX отчёта
    # -----------------------------

    doc = Document()

    doc.add_heading("AI QA Report", level=1)

    doc.add_heading("1. Анализ требований", level=2)
    doc.add_paragraph(analysis)

    doc.add_heading("2. Риски", level=2)
    doc.add_paragraph(risks)

    doc.add_heading("3. Тест-дизайн", level=2)
    doc.add_paragraph(test_design)

    doc.add_heading("4. Тест-кейсы", level=2)
    doc.add_paragraph(testcases)

    from io import BytesIO

    buffer = BytesIO()
    doc.save(buffer)

    st.download_button(
        label="Скачать QA отчет",
        data=buffer.getvalue(),
        file_name="qa_report.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

# -----------------------------
# AI QA CHAT
# -----------------------------

if st.session_state.last_result:

    st.divider()
    st.subheader("AI QA Chat — правки и уточнения")

    user_message = st.chat_input(
        "Например: добавь edge cases или перепиши тест-кейсы"
    )

    if user_message:

        result = st.session_state.last_result

        context = f"""
Ты Senior QA Engineer.

Текущий анализ требований:

{result.get("analysis","")}

Риски:

{result.get("risks","")}

Тест-дизайн:

{result.get("test_design","")}

Тест-кейсы:

{result.get("testcases","")}

Запрос пользователя:

{user_message}

Задача:

Исправь или дополни результат.

Можно:

- улучшить анализ
- добавить риски
- добавить тест-кейсы
- исправить тест-дизайн

Отвечай только на русском.
"""

        response = ""

        with st.chat_message("assistant"):

            message_placeholder = st.empty()

            for chunk in chat_llm.stream(context):

                response += chunk
                message_placeholder.markdown(response)

        st.session_state.chat_history.append({
            "role": "user",
            "content": user_message
        })

        st.session_state.chat_history.append({
            "role": "assistant",
            "content": response
        })


# -----------------------------
# отображение истории чата
# -----------------------------

for msg in st.session_state.chat_history:

    with st.chat_message(msg["role"]):
        st.write(msg["content"])