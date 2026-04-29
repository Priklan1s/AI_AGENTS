from graph.workflow import graph
import os
from openpyxl import Workbook
from docx import Document
from pypdf import PdfReader
from langgraph.graph import StateGraph
from agents.requirement_agent import requirement_agent

def read_docx(path):
    doc = Document(path)
    text = []

    for p in doc.paragraphs:
        text.append(p.text)

    return "\n".join(text)


def read_pdf(path):
    reader = PdfReader(path)
    text = []

    for page in reader.pages:
        text.append(page.extract_text())

    return "\n".join(text)


def load_requirements():

    print("1 — Ввести требования")
    print("2 — Загрузить файл")

    choice = input("Выбор: ")

    if choice == "1":

        print("Введите требования. Пустая строка — завершение")

        lines = []

        while True:
            line = input()

            if line == "":
                break

            lines.append(line)

        return "\n".join(lines)

    if choice == "2":

        path = input("Путь к файлу: ")

        if path.endswith(".txt"):
            with open(path, "r", encoding="utf-8") as f:
                return f.read()

        if path.endswith(".docx"):
            return read_docx(path)

        if path.endswith(".pdf"):
            return read_pdf(path)

        print("Неизвестный формат файла")
        return ""


def export_excel(csv_text):

    rows = [r.split(",") for r in csv_text.strip().split("\n")]

    wb = Workbook()
    ws = wb.active

    for row in rows:
        ws.append(row)

    os.makedirs("output", exist_ok=True)

    wb.save("output/testcases.xlsx")

    print("\nExcel файл создан: output/testcases.xlsx")


def export_word(text):

    doc = Document()

    doc.add_heading("Анализ требований", level=1)

    for line in text.split("\n"):
        doc.add_paragraph(line)

    os.makedirs("output", exist_ok=True)

    doc.save("output/analysis.docx")

    print("Word файл создан: output/analysis.docx")


def main():
    graph = StateGraph(dict)

    graph.add_node("analyze_requirements", requirement_agent)

    graph.set_entry_point("analyze_requirements")

    app = graph.compile()

    requirements = "Текст требований..."

    result = app.invoke({
        "requirements": requirements
    })

    print(result)

if __name__ == "__main__":
    main()